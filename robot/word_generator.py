import logging
import re
from datetime import date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import _Cell

from excel_parser import TeacherWorkload


DEPARTMENT_FULL_NAMES = {
    "КИ": "Компьютерная инженерия",
    "МКМ": "Медиа коммуникации и менеджмент",
    "ИИ": "Информатика и информационные технологии",
}


class WordGenerator:
    """Генератор индивидуального плана преподавателя по шаблону F-28_I-05.

    Шаблон — реальный бланк ХАТУ АҚ с 10 таблицами на kk/ru/en.
    Стратегия:
    1) Личные данные (Таблицы 0, 1) — positional replacement конкретных ячеек.
    2) Учебная нагрузка по семестрам (Таблица 2) — пересчёт сумм часов по
       строкам Excel и заполнение R2 (1 сем), R3 (2 сем), R4 (год).
    3) Финальная сводная (Таблица 9) — обновление колонок «План» в строках
       Учебная аудиторная/внеаудиторная (R2, R3).
    4) Учебный год — замена "20__-20__" в параграфах и ячейках на actual.
    5) Таблицы 3-8 (методическая/научная/организационная/воспитательная/
       повышение квалификации) — оставляем как есть в шаблоне; они
       заполняются вручную преподавателем.
    """

    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger or logging.getLogger("robot")

    def generate(
        self,
        workload: TeacherWorkload,
        template_path: Path,
        output_path: Path,
        academic_year: Optional[str] = None,
    ) -> Path:
        if not template_path.exists():
            raise FileNotFoundError(
                f"Word template not found: {template_path}. "
                "Положите шаблон F-28_I-05_*.docx как "
                f"{template_path.name} в папку templates/."
            )

        doc = Document(template_path)
        academic_year = academic_year or self._current_academic_year()
        year_start, year_end = academic_year.split("-")

        self._fill_personal_info(doc, workload)
        sem_summary = self._fill_workload_summary(doc, workload)
        self._fill_final_summary(doc, sem_summary)
        self._replace_academic_year(doc, year_start, year_end)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        self.logger.info(
            "Generated %s (teacher=%s, total=%.2fh)",
            output_path.name, workload.teacher_fio, workload.total_hours_year,
        )
        return output_path

    # --- Personal info -----------------------------------------------------

    def _fill_personal_info(self, doc, workload: TeacherWorkload) -> None:
        payment = self._humanize_payment_form(workload.payment_form)
        department = DEPARTMENT_FULL_NAMES.get(
            workload.op_department, workload.op_department
        )
        position = workload.position.strip()

        # Table 0 — главная шапка
        if len(doc.tables) > 0:
            t0 = doc.tables[0]
            if len(t0.rows) >= 6:
                # R4: тип занятости + должность (PhD остаётся как в шаблоне)
                self._set_cell(t0.rows[4], 1, payment)
                self._set_cell(t0.rows[4], 4, position)
                # R5: ФИО + кафедра
                self._set_cell(t0.rows[5], 1, workload.teacher_fio)
                self._set_cell(t0.rows[5], 4, department)

        # Table 1 — дубликат шапки внутри документа
        if len(doc.tables) > 1:
            t1 = doc.tables[1]
            if len(t1.rows) >= 3:
                self._set_cell(t1.rows[1], 1, payment)
                self._set_cell(t1.rows[1], 3, position)
                self._set_cell(t1.rows[2], 1, workload.teacher_fio)
                self._set_cell(t1.rows[2], 3, department)

    # --- Workload by semesters --------------------------------------------

    def _fill_workload_summary(
        self, doc, workload: TeacherWorkload
    ) -> dict:
        """Считает суммы часов по семестрам и заполняет Таблицу 2.

        Семестр строки определяется по тому, какая колонка
        (semester_1 или semester_2) имеет ненулевое значение в Excel.
        Если у строки заполнены обе — часы делятся между семестрами
        пропорционально (редкий случай для целогодовых дисциплин).
        Returns dict с агрегатами для Таблицы 9.
        """

        # Распределение строк по семестрам:
        # - sem1 only → весь час идёт в 1 сем
        # - sem2 only → весь час идёт во 2 сем
        # - оба заполнены ИЛИ оба пустые (year-long, напр. Руководство МД)
        #   → распределяем поровну, чтобы year == sem1 + sem2
        def in_sem1(r):
            return r.semester_1 > 0 and r.semester_2 == 0

        def in_sem2(r):
            return r.semester_2 > 0 and r.semester_1 == 0

        def shared(r):
            return (r.semester_1 > 0 and r.semester_2 > 0) or (
                r.semester_1 == 0 and r.semester_2 == 0
            )

        def sum_attr(rows, attr):
            return sum(getattr(r, attr) for r in rows)

        rows1 = [r for r in workload.load_rows if in_sem1(r)]
        rows2 = [r for r in workload.load_rows if in_sem2(r)]
        rows_half = [r for r in workload.load_rows if shared(r)]

        def s1(attr):
            return sum_attr(rows1, attr) + sum_attr(rows_half, attr) / 2

        def s2(attr):
            return sum_attr(rows2, attr) + sum_attr(rows_half, attr) / 2

        sem1 = {
            "lec": s1("lecture_hours"),
            "prac": s1("practical_hours"),
            "lab": s1("lab_hours"),
            "srsp": s1("srsp_hours"),
            "rk": s1("rk_hours"),
            "exam": s1("exam_hours"),
            # aud_total = ИТОГО из Excel (учитывает поточную логику и
            # коэффициенты), а не простая сумма компонент — иначе получим
            # инфляцию часов из-за повторяющихся групп в одном потоке.
            "aud_total": s1("total_hours"),
        }

        sem2 = {
            "lec": s2("lecture_hours"),
            "prac": s2("practical_hours"),
            "lab": s2("lab_hours"),
            "srsp": s2("srsp_hours"),
            "rk": s2("rk_hours"),
            "exam": s2("exam_hours"),
            "aud_total": s2("total_hours"),
        }

        year = {k: sem1[k] + sem2[k] for k in sem1}

        if len(doc.tables) >= 3:
            t = doc.tables[2]
            if len(t.rows) >= 5:
                self._fill_workload_row(t.rows[2], sem1)
                self._fill_workload_row(t.rows[3], sem2)
                self._fill_workload_row(t.rows[4], year)

        return {"sem1": sem1, "sem2": sem2, "year": year}

    def _fill_workload_row(self, row, data: dict) -> None:
        """Заполняет одну строку Таблицы 2 (14 колонок).

        Колонки (1-based):
        1=Лекции, 2=Практ, 3=Лаб, 4=СРОП/СРСП, 5=РК, 6=Экзамен,
        7=Всего ауд, 8-11=внеаудиторная (нули — нет в Excel), 12=Всего внеауд,
        13=Всего часов (== Всего ауд для нашего случая).
        """
        self._set_cell(row, 1, self._fmt(data["lec"]))
        self._set_cell(row, 2, self._fmt(data["prac"]))
        self._set_cell(row, 3, self._fmt(data["lab"]))
        self._set_cell(row, 4, self._fmt(data["srsp"]))
        self._set_cell(row, 5, self._fmt(data["rk"]))
        self._set_cell(row, 6, self._fmt(data["exam"]))
        self._set_cell(row, 7, self._fmt(data["aud_total"]))
        self._set_cell(row, 8, "0")
        self._set_cell(row, 9, "0")
        self._set_cell(row, 10, "0")
        self._set_cell(row, 11, "0")
        self._set_cell(row, 12, "0")
        self._set_cell(row, 13, self._fmt(data["aud_total"]))

    # --- Final summary (Table 9) ------------------------------------------

    def _fill_final_summary(self, doc, sem_summary: dict) -> None:
        """Обновляет колонки 'План' в финальной сводной (Таблица 9).

        Структура: C3=1сем план, C5=2сем план, C7=год план.
        R2: Учебная — Аудиторная (из Таблицы 2).
        R3: Учебная — Внеаудиторная (0 — нет в Excel).
        R4-R10 (методическая/научная/орг./воспитательная/повыш. квалификации)
        — оставляем как есть в шаблоне (заполняются вручную).
        """
        if len(doc.tables) < 10:
            return
        t = doc.tables[9]
        if len(t.rows) < 4:
            return

        sem1_aud = sem_summary["sem1"]["aud_total"]
        sem2_aud = sem_summary["sem2"]["aud_total"]
        year_aud = sem_summary["year"]["aud_total"]

        # R2: Аудиторная
        self._set_cell(t.rows[2], 3, self._fmt(sem1_aud))
        self._set_cell(t.rows[2], 5, self._fmt(sem2_aud))
        self._set_cell(t.rows[2], 7, self._fmt(year_aud))

        # R3: Внеаудиторная — обнуляем (нет данных в Excel)
        self._set_cell(t.rows[3], 3, "0")
        self._set_cell(t.rows[3], 5, "0")
        self._set_cell(t.rows[3], 7, "0")

    # --- Academic year ----------------------------------------------------

    def _replace_academic_year(self, doc, year_start: str, year_end: str) -> None:
        pattern = re.compile(r"20__-20__|20__-20__\s|20__\s*-\s*20__")
        replacement = f"{year_start}-{year_end}"

        def _replace_in_para(para):
            if pattern.search(para.text):
                # Простая замена через объединение runs (теряет inline-формат
                # отдельных runs, но сохраняет style параграфа).
                full = pattern.sub(replacement, para.text)
                for r in para.runs:
                    r.text = ""
                if para.runs:
                    para.runs[0].text = full
                else:
                    para.text = full

        for para in doc.paragraphs:
            _replace_in_para(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_para(para)

    # --- Helpers ----------------------------------------------------------

    @staticmethod
    def _humanize_payment_form(form: str) -> str:
        f = (form or "").strip().lower()
        if "совм" in f:
            return "совместитель"
        if "почас" in f:
            return "почасовик"
        if "штат" in f:
            return "штатный"
        return form or ""

    @staticmethod
    def _fmt(value: float) -> str:
        if value == int(value):
            return str(int(value))
        return f"{value:.2f}".rstrip("0").rstrip(".")

    @staticmethod
    def _set_cell(row, col_idx: int, value: str) -> None:
        """Безопасная установка значения ячейки с сохранением формата.

        Учитывает merged cells: если row.cells[col_idx] и row.cells[col_idx-1]
        указывают на одну ячейку (одинаковые _tc), не дублируем запись.
        Заменяет текст первого run, остальные runs параграфа очищаем.
        """
        if col_idx >= len(row.cells):
            return
        cell: _Cell = row.cells[col_idx]
        # Пропускаем если этот же _tc уже встречался слева (merged)
        for prev in range(col_idx):
            if row.cells[prev]._tc is cell._tc:
                return
        if not cell.paragraphs:
            cell.text = value
            return
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = value
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.text = value
        # Очищаем дополнительные параграфы (если в ячейке их несколько)
        for extra in cell.paragraphs[1:]:
            for r in extra.runs:
                r.text = ""

    @staticmethod
    def _current_academic_year() -> str:
        today = date.today()
        if today.month >= 7:
            return f"{today.year}-{today.year + 1}"
        return f"{today.year - 1}-{today.year}"
